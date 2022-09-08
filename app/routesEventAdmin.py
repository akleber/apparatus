from app import app, get_db, qrcode, utils
from flask import render_template, abort, url_for, send_file, request, redirect
from datetime import datetime
from pyexcel_xlsx import save_data
from collections import OrderedDict
from docx import Document
from htmldocx import HtmlToDocx
import markdown
from io import BytesIO
import uuid
import shortuuid


def get_event_data_verify_admin(adminToken: uuid, eventID: uuid) -> dict:
    cur = get_db().execute("SELECT * FROM event WHERE eventID = ?", (str(eventID),))
    rv = cur.fetchone()
    if not rv:
        app.logger.error(f"get_event_data_verify_admin: eventID unknown")
        abort(404)
    event_data = dict(rv)
    if event_data["adminToken"] != str(adminToken):
        app.logger.error(f"get_event_data_verify_admin: adminToke not as expected")
        abort(401)

    return event_data


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>", methods=["GET"])
def eventAdmin(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    activity_data = []
    cur = get_db().execute(
        "SELECT * FROM activity WHERE eventID = ? ORDER BY title", (str(eventID),)
    )
    for row in cur:
        activity_data.append(dict(row))

    return render_template(
        "eventAdmin.html", event_data=event_data, activity_data=activity_data
    )


@app.route("/eventAdmin/add", methods=["GET"])
def eventAdmin_event_add():
    event_data = {"eventID": str(uuid.uuid4()), "title": "", "description": ""}
    return render_template("eventEdit.html", event_data=event_data, add=True)


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>/edit", methods=["GET"])
def eventAdmin_event_edit(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    if event_data["adminToken"] != str(adminToken):
        app.logger.error(f"eventAdmin_event_edit: adminToken not as expected")
        return abort(401)

    return render_template("eventEdit.html", event_data=event_data, add=False)


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>/save", methods=["POST"])
def eventAdmin_event_save(adminToken, eventID):
    now = datetime.utcnow()

    blob = None
    if "bannerImage" in request.files:
        file = request.files["bannerImage"]
        if file.filename != "":
            blob = file.stream.read()

    cur = get_db().execute(
        "SELECT * FROM event WHERE eventID = ?",
        (str(eventID),),
    )
    rv = cur.fetchone()
    if rv:
        # update event
        event_data = dict(rv)
        if event_data["adminToken"] != str(adminToken):
            app.logger.error(f"eventAdmin_event_save: adminToken not as expected")
            abort(401)

        sql_data = {
            "eventID": str(eventID),
            "active": request.form.get("active", "0"),
            "title": request.form.get("title"),
            "description": request.form.get("description"),
            "legal": request.form.get("legal"),
            "lastChangedDate": now.isoformat(" "),
        }
        sql = """UPDATE event SET active = :active, title = :title, description = :description, legal = :legal, lastChangedDate = :lastChangedDate 
                 WHERE eventID = :eventID;"""

        if blob:
            sql_image = """UPDATE event SET bannerImage = ? WHERE eventID = ?;"""
            get_db().execute(sql_image, (memoryview(blob), str(eventID)))

    else:
        # add event
        userID, user_data = utils.add_user(
            request.form.get("firstName"),
            request.form.get("familyName"),
            request.form.get("mail"),
        )

        tinylink = shortuuid.uuid()[:10]

        sql_data = {
            "eventID": str(eventID),
            "tinylink": tinylink,
            "active": request.form.get("active", "0"),
            "title": request.form.get("title"),
            "description": request.form.get("description"),
            "legal": request.form.get("legal"),
            "creator": userID,
            "creationDate": now.isoformat(" "),
            "lastChangedDate": now.isoformat(" "),
            "adminToken": str(uuid.uuid4()),
            "bannerImage": blob,
        }
        sql = """INSERT INTO event (eventID, tinylink, active, title, description, legal, creator, creationDate, lastChangedDate, adminToken, bannerImage) 
                 VALUES (:eventID, :tinylink, :active, :title, :description, :legal, :creator, :creationDate, :lastChangedDate, :adminToken, :bannerImage);"""

    get_db().execute(sql, sql_data)
    get_db().commit()

    return redirect(url_for("eventAdmin", adminToken=adminToken, eventID=eventID))


@app.route("/eventAdmin/<uuid:eventID>/qr", methods=["GET"])
def qr(eventID):
    cur = get_db().execute(
        "SELECT tinylink FROM event WHERE eventID = ?", (str(eventID),)
    )
    rv = cur.fetchone()
    if not rv:
        app.logger.error(f"qr: eventID unknown")
        return abort(404)

    url = url_for("q", tinylink=rv["tinylink"])
    return send_file(
        qrcode(request.url_root[:-1] + url, mode="raw"), mimetype="image/png"
    )


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>/activity/add", methods=["GET"])
def eventAdmin_activity_add(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    activity_data = {"activityID": str(uuid.uuid4()), "title": "", "description": ""}
    return render_template(
        "activityEdit.html", event_data=event_data, activity_data=activity_data
    )


@app.route(
    "/eventAdmin/<uuid:adminToken>/<uuid:eventID>/activity/edit/<uuid:activityID>",
    methods=["GET"],
)
def eventAdmin_activity_edit(adminToken, eventID, activityID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    cur = get_db().execute(
        "SELECT * FROM activity WHERE eventID = ? AND activityID = ?",
        (
            str(eventID),
            str(activityID),
        ),
    )
    rv = cur.fetchone()
    if not rv:
        app.logger.error(f"eventAdmin_activity_edit: activity not found")
        return abort(404)
    activity_data = dict(rv)

    return render_template(
        "activityEdit.html", event_data=event_data, activity_data=activity_data
    )


@app.route(
    "/eventAdmin/<uuid:adminToken>/<uuid:eventID>/activity/save/<uuid:activityID>",
    methods=["POST"],
)
def eventAdmin_activity_save(adminToken, eventID, activityID):
    event_data = get_event_data_verify_admin(adminToken, eventID)
    now = datetime.utcnow()

    cur = get_db().execute(
        "SELECT * FROM activity WHERE activityID = ? and eventID = ?",
        (str(activityID), str(eventID)),
    )
    rv = cur.fetchone()
    if rv:
        user_data = {
            "activityID": str(activityID),
            "eventID": str(eventID),
            "active": request.form.get("active", "0"),
            "title": request.form.get("title"),
            "description": request.form.get("description"),
            "seats": request.form.get("seats", ""),
            "lastChangedDate": now.isoformat(" "),
        }
        sql = """UPDATE activity SET active = :active, title = :title, description = :description, seats = :seats, lastChangedDate = :lastChangedDate 
                 WHERE eventID = :eventID AND activityID = :activityID;"""

    else:
        user_data = {
            "activityID": str(activityID),
            "eventID": str(eventID),
            "active": request.form.get("active", "0"),
            "title": request.form.get("title"),
            "description": request.form.get("description"),
            "seats": request.form.get("seats", ""),
            "creationDate": now.isoformat(" "),
            "lastChangedDate": now.isoformat(" "),
        }
        sql = """INSERT INTO activity (activityID, eventID, active, title, description, seats, creationDate, lastChangedDate) 
                 VALUES (:activityID, :eventID, :active, :title, :description, :seats, :creationDate, :lastChangedDate);"""

    get_db().execute(sql, user_data)
    get_db().commit()

    return redirect(url_for("eventAdmin", adminToken=adminToken, eventID=eventID))


@app.route(
    "/eventAdmin/<uuid:adminToken>/<uuid:eventID>/activity/delete/<uuid:activityID>",
    methods=["GET"],
)
def eventAdmin_activity_delete(adminToken, eventID, activityID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    get_db().execute(
        "DELETE FROM activity WHERE activityID = ? and eventID = ?",
        (str(activityID), str(eventID)),
    )
    get_db().commit()

    return redirect(url_for("eventAdmin", adminToken=adminToken, eventID=eventID))


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>/attendees/list")
def eventAdmin_attendees_list(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    attendee_data = []
    cur = get_db().execute(
        "SELECT * FROM eventAttendees WHERE eventID = ? ORDER BY familyName",
        (str(eventID),),
    )
    for row in cur:
        r = dict(row)
        ags = r["AGs"]

        # TODO workaround as log as we do not have subtitles (for date, class)
        # We use the existing emojies as separator for the title and the subtitle
        cleanedAgs = ""
        skipping = False
        for i in ags:
            if i == "|":
                skipping = False
                cleanedAgs += i
            elif i >= "\U00001000":
                skipping = True
            elif skipping:
                continue
            else:
                cleanedAgs += i

        r["AGs"] = cleanedAgs

        attendee_data.append(r)

    return render_template(
        "attendeesList.html", event_data=event_data, attendee_data=attendee_data
    )


@app.route(
    "/eventAdmin/<uuid:adminToken>/<uuid:eventID>/attendees/delete/<uuid:attendeeID>"
)
def eventAdmin_attendees_delete(adminToken, eventID, attendeeID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    cur = get_db().execute(
        "SELECT * FROM attendee WHERE attendeeID = ?", (str(attendeeID),)
    )
    rv = cur.fetchone()
    if not rv:
        app.logger.error(f"eventAdmin_attendees_delete: attendeeID invalid")
        abort(403)

    attendee_data = dict(rv)

    get_db().execute("DELETE FROM user WHERE userID = ?", (attendee_data["userID"],))
    get_db().execute("DELETE FROM attendee WHERE attendeeID = ?", (str(attendeeID),))
    get_db().commit()

    return redirect(
        url_for(
            "eventAdmin_attendees_list",
            adminToken=str(adminToken),
            eventID=str(eventID),
        )
    )


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>/attendees/xlsx")
def eventAdmin_attendees_xlsx(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    excel_rows = []
    cur = get_db().execute(
        "SELECT * FROM eventAttendeesXlsx WHERE eventID = ?", (str(eventID),)
    )

    excel_col_names = []
    for col_name in cur.description:
        excel_col_names.append(col_name[0])
    excel_rows.append(excel_col_names)

    for row in cur:
        # TODO workaround remove everything after the emoji
        r = list(row)
        ag = r[11]
        cleanedAg = ""
        for i in ag:
            if i >= "\U00001000":
                break
            else:
                cleanedAg += i
        r[11] = cleanedAg

        excel_rows.append(r)

    data = OrderedDict()
    data.update({"Sheet 1": excel_rows})
    # data.update({"Sheet 2": [["row 1", "row 2", "row 3"]]})

    io = BytesIO()
    save_data(io, data)
    io.seek(0)

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"event_export_{timestamp}.xlsx"

    return send_file(io, as_attachment=True, download_name=filename, cache_timeout=0)


@app.route(
    "/eventAdmin/<uuid:adminToken>/<uuid:eventID>/activity/docx", methods=["GET"]
)
def eventAdmin_activity_docx(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    cur = get_db().execute(
        "SELECT title, description FROM activity WHERE eventID = ?", (str(eventID),)
    )

    document = Document()
    document.add_heading(event_data["title"], 0)

    new_parser = HtmlToDocx()

    for activity in cur:
        document.add_heading(activity["title"], level=1)

        md = activity["description"]
        html = markdown.markdown(md)

        new_parser.add_html_to_document(html, document)

    file = BytesIO()
    document.save(file)
    file.seek(0)

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"activity_export_{timestamp}.docx"

    return send_file(file, as_attachment=True, download_name=filename, cache_timeout=0)


@app.route("/eventAdmin/<uuid:adminToken>/<uuid:eventID>/duplicate", methods=["GET"])
def eventAdmin_duplicate(adminToken, eventID):
    event_data = get_event_data_verify_admin(adminToken, eventID)

    # make sure we have a verified user
    cur = get_db().execute(
        "SELECT * FROM user WHERE userID = ?", (event_data["creator"],)
    )
    rv = cur.fetchone()
    if not rv:
        app.logger.error(f"eventAdmin_duplicate: creator user not found")
        abort(500)
    user_data = dict(rv)
    if user_data["mailVerificationToken"]:
        app.logger.error(f"eventAdmin_duplicate: creator user not verified")
        abort(500)

    # duplicate user
    sql_data = {
        "firstName": user_data["firstName"],
        "familyName": user_data["familyName"],
        "mail": user_data["mail"],
        "mailVerificationToken": "",
        "gdprToken": str(uuid.uuid4()),
    }
    sql = """INSERT INTO user (firstName, familyName, mail, mailVerificationToken, gdprToken) 
             VALUES (:firstName, :familyName, :mail, :mailVerificationToken, :gdprToken);"""
    get_db().execute(sql, sql_data)

    now = datetime.utcnow()
    new_eventID = str(uuid.uuid4())
    new_adminToken = str(uuid.uuid4())

    # add duplicated event
    sql_data = {
        "eventID": new_eventID,
        "tinylink": shortuuid.uuid()[:10],
        "active": "0",
        "title": event_data["title"],
        "description": event_data["description"],
        "legal": event_data["legal"],
        "creator": event_data["creator"],
        "creationDate": now.isoformat(" "),
        "lastChangedDate": now.isoformat(" "),
        "adminToken": new_adminToken,
        "bannerImage": event_data["bannerImage"],
    }
    sql = """INSERT INTO event (eventID, tinylink, active, title, description, legal, creator, creationDate, lastChangedDate, adminToken, bannerImage) 
             VALUES (:eventID, :tinylink, :active, :title, :description, :legal, :creator, :creationDate, :lastChangedDate, :adminToken, :bannerImage);"""
    get_db().execute(sql, sql_data)

    # add duplicated activies
    activity_data = []
    cur = get_db().execute("SELECT * FROM activity WHERE eventID = ?", (str(eventID),))
    for row in cur:
        activity_data.append(dict(row))

    for activity in activity_data:
        sql_data = {
            "activityID": str(uuid.uuid4()),
            "eventID": new_eventID,
            "active": activity["active"],
            "title": activity["title"],
            "description": activity["description"],
            "seats": activity["seats"],
            "creationDate": now.isoformat(" "),
            "lastChangedDate": now.isoformat(" "),
        }
        sql = """INSERT INTO activity (activityID, eventID, active, title, description, seats, creationDate, lastChangedDate) 
                 VALUES (:activityID, :eventID, :active, :title, :description, :seats, :creationDate, :lastChangedDate);"""

        get_db().execute(sql, sql_data)

    get_db().commit()

    return redirect(
        url_for("eventAdmin", adminToken=new_adminToken, eventID=new_eventID)
    )
